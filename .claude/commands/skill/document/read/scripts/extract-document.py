#!/usr/bin/env python3
"""Extract text and images from PDF and Word documents."""

import base64
import sys
from pathlib import Path
from typing import Tuple


def extract_pdf(file_path: Path) -> Tuple[str, int, str]:
    """Extract text and images from PDF file using pymupdf.

    Args:
        file_path: Path to PDF file.

    Returns:
        Tuple of (content, page count, filename or error message).
    """
    try:
        import fitz
    except ImportError:
        return "", 0, "Error: pymupdf not installed. Run: pip install pymupdf"
    try:
        doc = fitz.open(str(file_path))
        page_count = len(doc)
        content = ""
        for idx in range(page_count):
            page = doc[idx]
            content += f"\n--- Page {idx + 1} ---\n"
            content += page.get_text()
            images = page.get_images(full=True)
            for img_idx, img in enumerate(images):
                xref = img[0]
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                img_ext = base_image["ext"]
                img_b64 = base64.b64encode(img_bytes).decode("utf-8")
                content += f"\n[IMAGE {img_idx + 1} — {img_ext.upper()} — base64]\n"
                content += f"data:image/{img_ext};base64,{img_b64}\n"
        doc.close()
        return content, page_count, file_path.name
    except Exception as e:
        return "", 0, f"Error reading PDF: {str(e)}"


def extract_docx(file_path: Path) -> Tuple[str, int, str]:
    """Extract text from Word document.

    Args:
        file_path: Path to DOCX file.

    Returns:
        Tuple of (text content, section count, filename or error message).
    """
    try:
        from docx import Document
    except ImportError:
        return "", 0, "Error: python-docx not installed. Run: pip install python-docx"
    try:
        doc = Document(str(file_path))
        text = ""
        section_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            section_count += 1
            text += "\n--- Table ---\n"
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                text += " | ".join(row_data) + "\n"
        section_count += len([p for p in doc.paragraphs if p.text.strip()])
        return text, section_count, file_path.name
    except Exception as e:
        return "", 0, f"Error reading DOCX: {str(e)}"


def main() -> None:
    """Main entry point for document extraction."""
    if len(sys.argv) < 2:
        print("Usage: python extract-document.py <file_path>")
        sys.exit(1)
    file_path = Path(sys.argv[1]).resolve()
    if not file_path.exists():
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        text, count, name = extract_pdf(file_path)
    elif suffix == ".docx":
        text, count, name = extract_docx(file_path)
    else:
        print(f"Error: Unsupported file format: {suffix}")
        print("Supported formats: .pdf, .docx")
        sys.exit(1)
    if not text:
        print(name)
        sys.exit(1)
    print(f"File: {name}")
    print(f"Type: {suffix.upper()}")
    print(f"Items: {count}")
    print(f"Characters: {len(text)}")
    print("\n" + "=" * 60 + "\n")
    print(text)


if __name__ == "__main__":
    main()
